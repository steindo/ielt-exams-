
import docx
import json
import re

def extract_unit_1(file_path):
    doc = docx.Document(file_path)
    content = []
    found_unit = False
    
    for para in doc.paragraphs:
        text = para.text.strip().replace("[TEXT FROM IMAGE]:", "").strip()
        if not text: continue
        
        # Detect Unit 1 start: "1  Getting ready to listen" or something similar
        if re.match(r'^1\s+Getting ready to listen', text):
            found_unit = True
            content.append(text)
            print(f"Found Unit 1: {text}")
            continue
            
        if found_unit:
            # Stop if we hit Unit 2: "2  Following a conversation"
            if re.match(r'^2\s+Following a conversation', text):
                print(f"Stopping at Unit 2: {text}")
                break
            content.append(text)
            
    return content

def extract_scripts(file_path):
    doc = docx.Document(file_path)
    scripts = []
    found_scripts = False
    
    for para in doc.paragraphs:
        text = para.text.strip().replace("[TEXT FROM IMAGE]:", "").strip()
        if "Recording Scripts" in text:
            found_scripts = True
            print("Found Recording Scripts section.")
            continue
        
        if found_scripts:
            scripts.append(text)
            
    return scripts

file_path = r'c:\Users\user\Desktop\malik\ielts apps\output_formatted.docx'
unit_1 = extract_unit_1(file_path)
with open("unit_1_full.json", "w", encoding="utf-8") as f:
    json.dump(unit_1, f, indent=2, ensure_ascii=False)

scripts = extract_scripts(file_path)
with open("recording_scripts_raw.json", "w", encoding="utf-8") as f:
    json.dump(scripts, f, indent=2, ensure_ascii=False)

print("Extraction complete.")
