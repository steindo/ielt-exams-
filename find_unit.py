from docx import Document

doc = Document(r'c:\Users\user\Desktop\malik\ielts apps\output_formatted.docx')
found_at = -1
for i, para in enumerate(doc.paragraphs):
    if "Getting ready to listen" in para.text and i > 50: # Skip the index
        found_at = i
        break

if found_at != -1:
    print(f"Unit 1 found at paragraph {found_at}")
    print("--- CONTENT PREVIEW ---")
    for j in range(found_at, found_at + 50):
        if j < len(doc.paragraphs):
            print(f"[{j}] {doc.paragraphs[j].text[:120]}")
else:
    print("Unit 1 title not found after the index.")
