from docx import Document
import sys

# Force UTF-8 output
sys.stdout.reconfigure(encoding='utf-8')

try:
    doc = Document(r'c:\Users\user\Desktop\malik\ielts apps\output_formatted.docx')
    
    print("--- SEARCHING FOR 'UNIT' HEADERS ---")
    
    unit_count = 0
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        # Heuristic: Unit titles are often short and start with "Unit"
        if text.lower().startswith("unit") and len(text) < 50:
            print(f"\n[POSSIBLE HEADER line {i}]: {text}")
            # Print next few lines to see the topic
            for j in range(1, 4):
                if i + j < len(doc.paragraphs):
                    print(f"   + {doc.paragraphs[i+j].text.strip()}")
            unit_count += 1
            
    print(f"\nFound {unit_count} possible units.")
    
    print("\n--- SAMPLE CONTENT (First 50 lines) ---")
    for i, para in enumerate(doc.paragraphs[:50]):
         print(f"{i}: {para.text.strip()}")

except Exception as e:
    print(f"Error: {e}")
